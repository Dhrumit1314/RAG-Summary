"""
Document Processor for Model Validation Work Plan Automation
"""
import asyncio
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
import re

try:
    import pypdf
    from pypdf import PdfReader
except ImportError:
    pypdf = None

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from config import config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes various document formats for model validation"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.md', '.docx', '.csv']
        self.max_content_length = 500000  # Increased for white papers (500K chars)
    
    async def process_document(self, file_path: str) -> str:
        """Process document and extract text content"""
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            if file_path_obj.suffix.lower() not in self.supported_formats:
                logger.warning(f"Unsupported file format: {file_path_obj.suffix}")
                return ""
            
            # Process based on file type
            if file_path_obj.suffix.lower() == '.pdf':
                return await self._process_pdf(file_path_obj)
            elif file_path_obj.suffix.lower() == '.docx':
                return await self._process_docx(file_path_obj)
            elif file_path_obj.suffix.lower() == '.txt':
                return await self._process_text(file_path_obj)
            elif file_path_obj.suffix.lower() == '.md':
                return await self._process_markdown(file_path_obj)
            elif file_path_obj.suffix.lower() == '.csv':
                return await self._process_csv(file_path_obj)
            else:
                return await self._process_generic(file_path_obj)
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return ""
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF documents"""
        if not pypdf:
            logger.warning("PyPDF not available for PDF processing")
            return ""
        
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    if page_num >= 150:  # Limit to first 150 pages for white paper
                        break
                    
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            return self._clean_content(content)
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return ""
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process Word documents (.docx)"""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available for Word document processing")
            return ""
        
        try:
            content = ""
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text.strip() + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content += cell.text.strip() + " | "
                    content += "\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                for header in section.header.paragraphs:
                    if header.text.strip():
                        content += f"[Header] {header.text.strip()}\n"
                for footer in section.footer.paragraphs:
                    if footer.text.strip():
                        content += f"[Footer] {footer.text.strip()}\n"
            
            return self._clean_content(content)
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            return ""
    
    async def _process_text(self, file_path: Path) -> str:
        """Process text documents"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return self._clean_content(content)
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return ""
    
    async def _process_markdown(self, file_path: Path) -> str:
        """Process markdown documents"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Remove markdown formatting for better text extraction
            content = re.sub(r'#+\s+', '', content)  # Remove headers
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Remove bold
            content = re.sub(r'\*(.*?)\*', r'\1', content)  # Remove italic
            content = re.sub(r'`(.*?)`', r'\1', content)  # Remove code
            content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)  # Remove links
            
            return self._clean_content(content)
            
        except Exception as e:
            logger.error(f"Error processing markdown file {file_path}: {e}")
            return ""
    
    async def _process_csv(self, file_path: Path) -> str:
        """Process CSV documents"""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            content = f"CSV Document: {file_path.name}\n"
            content += f"Shape: {df.shape}\n"
            content += f"Columns: {list(df.columns)}\n\n"
            
            # Add sample data
            if len(df) > 0:
                content += "Sample Data:\n"
                content += df.head(10).to_string()
            
            return self._clean_content(content)
            
        except ImportError:
            logger.warning("Pandas not available for CSV processing")
            return await self._process_text(file_path)
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            return ""
    
    async def _process_generic(self, file_path: Path) -> str:
        """Process generic documents"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            return self._clean_content(content)
        except Exception as e:
            logger.error(f"Error processing generic file {file_path}: {e}")
            return ""
    
    def _clean_content(self, content: str) -> str:
        """Clean and normalize document content"""
        if not content:
            return ""
        
        # Remove excessive whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove special characters that might interfere with LLM processing
        content = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]]', '', content)
        
        # Limit content length
        if len(content) > self.max_content_length:
            content = content[:self.max_content_length] + "... [Content truncated]"
        
        return content.strip()
    
    async def analyze_document(self, content: str, analysis_type: str = "general") -> Dict[str, Any]:
        """Analyze document content for validation insights"""
        try:
            # Basic content analysis
            analysis = {
                "analysis_type": analysis_type,
                "content_length": len(content),
                "word_count": len(content.split()),
                "key_topics": self._extract_key_topics(content),
                "risk_indicators": self._identify_risk_indicators(content),
                "compliance_keywords": self._identify_compliance_keywords(content),
                "technical_terms": self._identify_technical_terms(content),
                "summary": self._generate_content_summary(content)
            }
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing document: {e}")
            return {"error": str(e)}
    
    def _extract_key_topics(self, content: str) -> List[str]:
        """Extract key topics from document content"""
        # Simple keyword extraction (in production, use more sophisticated NLP)
        keywords = [
            "model", "validation", "risk", "compliance", "data", "performance",
            "assessment", "review", "testing", "documentation", "governance",
            "regulatory", "business", "logic", "assumptions", "methodology"
        ]
        
        found_topics = []
        content_lower = content.lower()
        
        for keyword in keywords:
            if keyword in content_lower:
                found_topics.append(keyword)
        
        return found_topics[:10]  # Limit to top 10 topics
    
    def _identify_risk_indicators(self, content: str) -> List[str]:
        """Identify risk-related indicators in content"""
        risk_keywords = [
            "risk", "uncertainty", "volatility", "exposure", "vulnerability",
            "threat", "hazard", "danger", "instability", "fluctuation"
        ]
        
        found_risks = []
        content_lower = content.lower()
        
        for keyword in risk_keywords:
            if keyword in content_lower:
                found_risks.append(keyword)
        
        return found_risks[:5]
    
    def _identify_compliance_keywords(self, content: str) -> List[str]:
        """Identify compliance-related keywords"""
        compliance_keywords = [
            "compliance", "regulation", "regulatory", "requirement", "standard",
            "guideline", "policy", "procedure", "audit", "oversight"
        ]
        
        found_compliance = []
        content_lower = content.lower()
        
        for keyword in compliance_keywords:
            if keyword in content_lower:
                found_compliance.append(keyword)
        
        return found_compliance[:5]
    
    def _identify_technical_terms(self, content: str) -> List[str]:
        """Identify technical terms and methodologies"""
        technical_terms = [
            "algorithm", "methodology", "framework", "architecture", "implementation",
            "parameter", "calibration", "backtesting", "stress testing", "scenario",
            "monte carlo", "simulation", "optimization", "machine learning", "AI"
        ]
        
        found_technical = []
        content_lower = content.lower()
        
        for term in technical_terms:
            if term in content_lower:
                found_technical.append(term)
        
        return found_technical[:8]
    
    def _generate_content_summary(self, content: str) -> str:
        """Generate a brief summary of the content"""
        if len(content) < 200:
            return content
        
        # Simple summary: first and last sentences
        sentences = content.split('.')
        if len(sentences) >= 2:
            summary = sentences[0].strip() + ". " + sentences[-1].strip()
            if len(summary) > 300:
                summary = summary[:300] + "..."
            return summary
        else:
            return content[:300] + "..." if len(content) > 300 else content
    
    async def extract_validation_requirements(self, content: str) -> Dict[str, Any]:
        """Extract validation requirements from document content"""
        try:
            requirements = {
                "data_requirements": self._extract_data_requirements(content),
                "performance_requirements": self._extract_performance_requirements(content),
                "compliance_requirements": self._extract_compliance_requirements(content),
                "documentation_requirements": self._extract_documentation_requirements(content),
                "timeline_requirements": self._extract_timeline_requirements(content)
            }
            
            return requirements
            
        except Exception as e:
            logger.error(f"Error extracting validation requirements: {e}")
            return {"error": str(e)}
    
    def _extract_data_requirements(self, content: str) -> List[str]:
        """Extract data-related requirements"""
        data_patterns = [
            r"data\s+quality", r"data\s+validation", r"data\s+integrity",
            r"data\s+governance", r"data\s+lineage", r"data\s+accuracy"
        ]
        
        requirements = []
        for pattern in data_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            requirements.extend(matches)
        
        return list(set(requirements))[:5]
    
    def _extract_performance_requirements(self, content: str) -> List[str]:
        """Extract performance-related requirements"""
        perf_patterns = [
            r"performance\s+metrics", r"accuracy\s+requirements", r"precision\s+targets",
            r"recall\s+standards", r"f1\s+score", r"auc\s+requirements"
        ]
        
        requirements = []
        for pattern in perf_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            requirements.extend(matches)
        
        return list(set(requirements))[:5]
    
    def _extract_compliance_requirements(self, content: str) -> List[str]:
        """Extract compliance-related requirements"""
        comp_patterns = [
            r"regulatory\s+compliance", r"audit\s+requirements", r"oversight\s+standards",
            r"governance\s+framework", r"risk\s+management", r"control\s+requirements"
        ]
        
        requirements = []
        for pattern in comp_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            requirements.extend(matches)
        
        return list(set(requirements))[:5]
    
    def _extract_documentation_requirements(self, content: str) -> List[str]:
        """Extract documentation-related requirements"""
        doc_patterns = [
            r"documentation\s+standards", r"model\s+documentation", r"technical\s+specifications",
            r"user\s+manuals", r"maintenance\s+procedures", r"change\s+management"
        ]
        
        requirements = []
        for pattern in doc_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            requirements.extend(matches)
        
        return list(set(requirements))[:5]
    
    def _extract_timeline_requirements(self, content: str) -> List[str]:
        """Extract timeline-related requirements"""
        time_patterns = [
            r"timeline", r"deadline", r"schedule", r"milestone", r"deliverable",
            r"quarterly", r"annual", r"monthly", r"weekly", r"daily"
        ]
        
        requirements = []
        for pattern in time_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            requirements.extend(matches)
        
        return list(set(requirements))[:5]
